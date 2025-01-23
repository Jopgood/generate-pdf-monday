import os
from docx import Document
from pydantic import BaseModel
from subprocess import run, PIPE
import tempfile
import re
from logging_config import LoggerAdapter, setup_logging
import logging

logger = LoggerAdapter(setup_logging(), "document_service")


class DocxEditOptions(BaseModel):
    input_file_name: str
    output_file_name: str
    placeholders: dict


def sanitize_filename(filename):
    # Remove any non-word characters (everything except numbers and letters)
    filename = re.sub(r"[^\w\s-]", "", filename)
    # Replace all runs of whitespace with a single underscore
    filename = re.sub(r"\s+", "_", filename)
    return filename


def convert_to_pdf(input_path: str, output_path: str):
    # Using LibreOffice directly for conversion
    result = run(['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir',
                  os.path.dirname(output_path), input_path],
                 stdout=PIPE, stderr=PIPE)
    if result.returncode != 0:
        raise Exception(f"PDF conversion failed: {result.stderr.decode()}")

    # Rename the output file to match the desired name
    temp_pdf_path = os.path.join(os.path.dirname(output_path),
                                 os.path.basename(input_path).rsplit('.', 1)[0] + '.pdf')
    os.rename(temp_pdf_path, output_path)


def edit_docx_convert_to_pdf_and_save(options: DocxEditOptions) -> str:
    logger.log_operation(
        logging.INFO,
        "edit_document",
        "Starting document editing process",
        {"input_file": options.input_file_name,
            "output_file": options.output_file_name}
    )

    project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    templates_dir = os.path.join(project_root, 'templates')
    output_dir = os.path.join(project_root, 'output')

    input_docx_path = os.path.join(templates_dir, options.input_file_name)
    output_pdf_path = os.path.join(
        output_dir, f"{sanitize_filename(options.output_file_name)}.pdf")

    try:
        os.makedirs(output_dir, exist_ok=True)
        doc = Document(input_docx_path)

        logger.log_operation(
            logging.INFO,
            "edit_document",
            "Processing document placeholders"
        )
        os.makedirs(output_dir, exist_ok=True)

        doc = Document(input_docx_path)

        for paragraph in doc.paragraphs:
            for key, value in options.placeholders.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, value)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in options.placeholders.items():
                            if key in paragraph.text:
                                paragraph.text = paragraph.text.replace(
                                    key, value)

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            temp_docx_path = tmp.name
            doc.save(temp_docx_path)

        logger.log_operation(
            logging.INFO,
            "convert_pdf",
            "Converting document to PDF",
            {"temp_path": temp_docx_path}
        )

        convert_to_pdf(temp_docx_path, output_pdf_path)

        logger.log_operation(
            logging.INFO,
            "edit_document",
            "Document processing completed successfully",
            {"output_path": output_pdf_path}
        )

        return output_pdf_path
    except Exception as e:
        logger.log_operation(
            logging.ERROR,
            "edit_document",
            "Document processing failed",
            {"error": str(e)},
            exc_info=True
        )
        raise

    finally:
        if 'temp_docx_path' in locals() and os.path.exists(temp_docx_path):
            os.remove(temp_docx_path)
            logger.log_operation(
                logging.INFO,
                "cleanup",
                "Temporary file deleted",
                {"temp_path": temp_docx_path}
            )
